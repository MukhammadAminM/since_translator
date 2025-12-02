"""
Парсер для извлечения терминов из глоссария
Поддерживает форматы: TXT, PDF, DOCX
С поддержкой OCR для PDF с изображениями
"""
import re
import json
import os
from pathlib import Path
from typing import Dict, List, Tuple, Optional

# OCR для распознавания текста из изображений
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
    # Настройка пути к Tesseract для Windows (если нужно)
    if os.name == 'nt':  # Windows
        tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        for path in tesseract_paths:
            if Path(path).exists():
                pytesseract.pytesseract.tesseract_cmd = path
                break
except ImportError:
    OCR_AVAILABLE = False

# Пробуем разные библиотеки для работы с PDF
try:
    import PyPDF2
    PDF_AVAILABLE = True
    PDF_LIB = "PyPDF2"
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
        PDF_LIB = "pdfplumber"
    except ImportError:
        try:
            import fitz  # PyMuPDF
            PDF_AVAILABLE = True
            PDF_LIB = "pymupdf"
        except ImportError:
            PDF_AVAILABLE = False
            PDF_LIB = None

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class GlossaryParser:
    """
    Парсер для извлечения терминов из файлов глоссария
    """
    
    def __init__(self, glossary_dir: str = "glossary"):
        self.glossary_dir = Path(glossary_dir)
    
    def parse_txt_file(self, file_path: Path) -> List[Tuple[str, str]]:
        """
        Парсит TXT файл с терминами в формате:
        "Русский термин – English translation"
        или
        "1. Русский термин – English translation"
        """
        terms = []
        
        with open(file_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                
                # Убираем нумерацию в начале строки
                line = re.sub(r'^\d+\.\s*', '', line)
                
                # Ищем разделитель "–" или "-"
                if "–" in line:
                    parts = line.split("–", 1)
                elif " - " in line:
                    parts = line.split(" - ", 1)
                elif " — " in line:
                    parts = line.split(" — ", 1)
                else:
                    continue
                
                if len(parts) == 2:
                    source_term = self._normalize_text(parts[0].strip())
                    target_term = self._normalize_text(parts[1].strip())
                    
                    if source_term and target_term:
                        terms.append((source_term, target_term))
        
        return terms
    
    def parse_pdf_file(self, file_path: Path) -> List[Tuple[str, str]]:
        """
        Извлекает текст из PDF и парсит термины
        Поддерживает PyPDF2, pdfplumber и PyMuPDF
        """
        if not PDF_AVAILABLE:
            raise ImportError(
                "PDF библиотека не установлена. Установите одну из: "
                "pip install PyPDF2 или pip install pdfplumber или pip install pymupdf"
            )
        
        text = ""
        
        # Извлекаем текст в зависимости от доступной библиотеки
        if PDF_LIB == "PyPDF2":
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        elif PDF_LIB == "pdfplumber":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif PDF_LIB == "pymupdf":
            import fitz
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        
        # Если текст не извлечен или слишком короткий, пробуем OCR (PDF может содержать изображения)
        if not text.strip() or len(text.strip()) < 50:
            if not text.strip():
                print(f"⚠️  Текст не извлечен из {file_path.name}, пробуем OCR...")
            else:
                print(f"⚠️  Мало текста из {file_path.name} ({len(text.strip())} символов), пробуем OCR...")
            ocr_text = self._extract_text_with_ocr(file_path)
            if ocr_text:
                text = ocr_text
        
        if not text.strip():
            print(f"❌ Не удалось извлечь текст из {file_path.name}")
            return []
        
        # Парсим текст аналогично TXT
        # Обрабатываем как однострочный, так и многострочный формат (для арабского)
        terms = []
        lines = [line.strip() for line in text.split("\n")]
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Пропускаем пустые строки
            if not line:
                i += 1
                continue
            
            # Сохраняем оригинальную строку для многострочного формата
            original_line = line
            
            # Убираем нумерацию
            line = re.sub(r'^\d+\.\s*', '', line)
            
            # Если после удаления нумерации строка пустая, переходим к следующей
            if not line:
                i += 1
                continue
            
            # Проверяем однострочный формат (термин – перевод)
            separators = ["–", "—", " - ", " -", "- ", ":", "=", "→"]
            parts = None
            
            for sep in separators:
                if sep in line:
                    parts = line.split(sep, 1)
                    break
            
            # Если не найден в строке, пробуем regex
            if not parts or len(parts) != 2:
                match = re.search(r'([\u0600-\u06FF\s]+)[\s\-–—:=→]+([A-Za-z\s\(\)]+)', line)
                if match:
                    parts = [match.group(1).strip(), match.group(2).strip()]
                else:
                    match = re.search(r'([A-Za-z\s\(\)]+)[\s\-–—:=→]+([\u0600-\u06FF\s]+)', line)
                    if match:
                        parts = [match.group(2).strip(), match.group(1).strip()]
            
            # Если не найден однострочный формат, пробуем многострочный
            # Формат: арабский_термин (может быть на нескольких строках)\n–\nанглийский_перевод
            if (not parts or len(parts) != 2) and i + 1 < len(lines):
                # Проверяем, есть ли арабский текст в текущей строке
                arabic_in_current = re.search(r'[\u0600-\u06FF]', line)
                
                if arabic_in_current:
                    # Собираем арабский термин из нескольких строк (до разделителя)
                    arabic_parts = [line]
                    j = i + 1
                    
                    # Ищем разделитель
                    while j < len(lines):
                        if not lines[j].strip():
                            j += 1
                            continue
                        # Если нашли разделитель, останавливаемся
                        if lines[j].strip() in ["–", "—", "-"]:
                            separator_line = j
                            j += 1
                            break
                        # Если встретили номер или английский текст, это не арабский термин
                        if re.match(r'^\d+\.', lines[j]) or re.search(r'[A-Za-z]', lines[j]):
                            break
                        # Если есть арабский текст, добавляем к термину
                        if re.search(r'[\u0600-\u06FF]', lines[j]):
                            arabic_parts.append(lines[j])
                        j += 1
                    
                    # Если нашли разделитель, собираем английский перевод
                    if 'separator_line' in locals() and j < len(lines):
                        source_term = " ".join(arabic_parts)
                        target_parts = []
                        
                        # Собираем английский перевод после разделителя
                        while j < len(lines):
                            if not lines[j].strip():
                                j += 1
                                continue
                            # Если встретили новый термин (арабский или номер), останавливаемся
                            if (re.search(r'[\u0600-\u06FF]', lines[j]) or 
                                re.match(r'^\d+\.', lines[j]) or
                                lines[j].strip() in ["–", "—", "-"]):
                                break
                            # Добавляем английский текст
                            if re.search(r'[A-Za-z]', lines[j]):
                                target_parts.append(lines[j].strip())
                            j += 1
                        
                        if target_parts:
                            target_term = " ".join(target_parts)
                            parts = [source_term, target_term]
                            i = j - 1  # Пропускаем обработанные строки
            
            if parts and len(parts) == 2:
                source_term = parts[0].strip()
                target_term = parts[1].strip()
                
                # Проверяем, что есть хотя бы один непустой термин
                if source_term and target_term:
                    # Убираем лишние пробелы
                    source_term = re.sub(r'\s+', ' ', source_term)
                    target_term = re.sub(r'\s+', ' ', target_term)
                    # Нормализуем кодировку
                    source_term = self._normalize_text(source_term)
                    target_term = self._normalize_text(target_term)
                    terms.append((source_term, target_term))
            
            i += 1
        
        return terms
    
    def parse_docx_file(self, file_path: Path) -> List[Tuple[str, str]]:
        """
        Извлекает текст из DOCX и парсит термины
        Обрабатывает как параграфы, так и таблицы
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx не установлен. Установите: pip install python-docx")
        
        doc = Document(file_path)
        terms = []
        
        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            line = paragraph.text.strip()
            if not line:
                continue
            
            # Убираем нумерацию
            line = re.sub(r'^\d+\.\s*', '', line)
            
            # Ищем разделитель
            if "–" in line:
                parts = line.split("–", 1)
            elif " - " in line:
                parts = line.split(" - ", 1)
            elif " — " in line:
                parts = line.split(" — ", 1)
            else:
                continue
            
            if len(parts) == 2:
                source_term = parts[0].strip()
                target_term = parts[1].strip()
                
                if source_term and target_term:
                    terms.append((source_term, target_term))
        
        # Обрабатываем таблицы (формат: аббревиатура | разделитель | расшифровка)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                
                # Пропускаем пустые строки и заголовки
                if not any(cells) or len(cells) < 2:
                    continue
                
                # Формат 1: [Аббревиатура, "-", Расшифровка] (3 колонки)
                if len(cells) >= 3:
                    abbrev = cells[0].strip()
                    separator = cells[1].strip()
                    expansion = cells[2].strip()
                    
                    # Проверяем, что средняя колонка - разделитель
                    if separator in ["-", "–", "—", ":", "="] and abbrev and expansion:
                        # Аббревиатура -> Расшифровка
                        abbrev = self._normalize_text(abbrev)
                        expansion = self._normalize_text(expansion)
                        terms.append((abbrev, expansion))
                        continue
                
                # Формат 2: [Аббревиатура, Расшифровка] (2 колонки)
                if len(cells) >= 2:
                    cell1 = cells[0].strip()
                    cell2 = cells[1].strip()
                    
                    # Если первая колонка короткая (вероятно аббревиатура) и вторая длинная
                    if (len(cell1) <= 20 and len(cell2) > len(cell1) and 
                        cell1 and cell2 and 
                        not re.search(r'[–—]', cell1)):  # Нет разделителя в первой колонке
                        cell1 = self._normalize_text(cell1)
                        cell2 = self._normalize_text(cell2)
                        terms.append((cell1, cell2))
                        continue
                
                # Формат 3: [Термин - Перевод] в одной ячейке
                for cell_text in cells:
                    if not cell_text:
                        continue
                    
                    # Ищем разделитель в ячейке
                    if "–" in cell_text:
                        parts = cell_text.split("–", 1)
                    elif " - " in cell_text:
                        parts = cell_text.split(" - ", 1)
                    elif " — " in cell_text:
                        parts = cell_text.split(" — ", 1)
                    else:
                        continue
                    
                    if len(parts) == 2:
                        source_term = self._normalize_text(parts[0].strip())
                        target_term = self._normalize_text(parts[1].strip())
                        
                        if source_term and target_term:
                            terms.append((source_term, target_term))
        
        return terms
    
    def _extract_text_with_ocr(self, file_path: Path) -> str:
        """
        Извлекает текст из PDF используя OCR (распознавание текста из изображений)
        Требует установленный Tesseract OCR и poppler
        """
        if not OCR_AVAILABLE:
            print("⚠️  OCR библиотеки не установлены. Установите: pip install pytesseract pdf2image Pillow")
            print("   Также требуется установить Tesseract OCR: https://github.com/tesseract-ocr/tesseract")
            return ""
        
        try:
            # Конвертируем PDF в изображения
            print(f"   Конвертация PDF в изображения...")
            
            # Пробуем найти poppler в стандартных местах
            poppler_path = None
            if os.name == 'nt':  # Windows
                poppler_paths = [
                    r'C:\poppler\Library\bin',
                    r'C:\poppler\bin',
                    r'C:\Program Files\poppler\bin',
                ]
                for path in poppler_paths:
                    if Path(path).exists():
                        poppler_path = path
                        break
            
            # Конвертируем PDF в изображения
            if poppler_path:
                images = convert_from_path(str(file_path), dpi=300, poppler_path=poppler_path)
            else:
                images = convert_from_path(str(file_path), dpi=300)
            
            text = ""
            print(f"   Распознавание текста из {len(images)} страниц...")
            
            for i, image in enumerate(images, 1):
                print(f"   Страница {i}/{len(images)}...", end="\r")
                # Распознаем текст с поддержкой русского языка
                page_text = pytesseract.image_to_string(
                    image, 
                    lang='rus+eng',  # Русский и английский
                    config='--psm 6'  # Предполагаем единый блок текста
                )
                text += page_text + "\n"
            
            print(f"   ✅ Распознано {len(text)} символов")
            return text
            
        except Exception as e:
            print(f"   ❌ Ошибка OCR: {str(e)}")
            print("   Убедитесь, что Tesseract OCR установлен и доступен в PATH")
            return ""
    
    def parse_excel_file(self, file_path: Path) -> List[Tuple[str, str]]:
        """
        Парсит Excel файл с терминами (предполагается формат: колонка A - исходный термин, колонка B - перевод)
        """
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl не установлен. Установите: pip install openpyxl")
        
        terms = []
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        
        for row in sheet.iter_rows(min_row=2, values_only=True):  # Пропускаем заголовок
            if row[0] and row[1]:  # Если обе ячейки заполнены
                source_term = str(row[0]).strip()
                target_term = str(row[1]).strip()
                
                if source_term and target_term:
                    # Нормализуем кодировку терминов
                    source_term = self._normalize_text(source_term)
                    target_term = self._normalize_text(target_term)
                    terms.append((source_term, target_term))
        
        return terms
    
    def _normalize_text(self, text: str) -> str:
        """
        Нормализует текст, исправляя проблемы с кодировкой
        """
        if not text:
            return text
        
        try:
            # Проверяем, есть ли подозрительные символы (результат неправильной кодировки)
            # Символы типа ɍ, ɓ, ɬ, ɨ и т.д. - это обычно результат неправильной интерпретации UTF-8
            has_suspicious = False
            for char in text:
                code = ord(char)
                # Символы в диапазоне 400-600, которые не являются кириллицей или латиницей
                # Это часто результат неправильной кодировки
                if (400 <= code <= 600 and 
                    not (1040 <= code <= 1103) and  # не кириллица
                    not (65 <= code <= 90) and      # не латиница верхний
                    not (97 <= code <= 122)):       # не латиница нижний
                    has_suspicious = True
                    break
            
            # Если найдены подозрительные символы, пробуем исправить
            if has_suspicious:
                # Пробуем разные варианты декодирования
                for encoding in ['cp1251', 'cp866', 'iso-8859-5']:
                    try:
                        # Конвертируем через latin1 (чтобы получить байты) и затем декодируем
                        text_bytes = text.encode('latin1', errors='ignore')
                        decoded = text_bytes.decode(encoding, errors='ignore')
                        # Проверяем, что результат лучше (меньше подозрительных символов)
                        suspicious_count = sum(1 for c in decoded if 400 <= ord(c) <= 600 and 
                                             not (1040 <= ord(c) <= 1103) and
                                             not (65 <= ord(c) <= 90) and
                                             not (97 <= ord(c) <= 122))
                        original_suspicious = sum(1 for c in text if 400 <= ord(c) <= 600 and 
                                                not (1040 <= ord(c) <= 1103) and
                                                not (65 <= ord(c) <= 90) and
                                                not (97 <= ord(c) <= 122))
                        if suspicious_count < original_suspicious:
                            text = decoded
                            break
                    except:
                        continue
            
            # Убеждаемся, что текст в UTF-8
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            
        except Exception as e:
            # Если не удалось исправить, возвращаем как есть
            pass
        
        return text
    
    def _has_valid_text(self, text: str) -> bool:
        """
        Проверяет, что текст содержит валидные символы (не только результат неправильной кодировки)
        """
        if not text:
            return False
        
        # Подсчитываем валидные символы (кириллица, латиница, арабские, китайские, цифры, знаки препинания)
        valid_count = 0
        suspicious_count = 0
        
        for char in text:
            code = ord(char)
            # Валидные символы
            if (1040 <= code <= 1103 or  # кириллица
                65 <= code <= 90 or      # латиница верхний
                97 <= code <= 122 or     # латиница нижний
                48 <= code <= 57 or       # цифры
                0x0600 <= code <= 0x06FF or  # арабский
                0x4E00 <= code <= 0x9FFF or  # китайский
                char in ".,;:!?()[]{}\"'/-–—=+*&%$#@ "):  # знаки препинания и пробелы
                valid_count += 1
            # Подозрительные символы (результат неправильной кодировки)
            elif 400 <= code <= 600:
                suspicious_count += 1
        
        # Если больше подозрительных символов, чем валидных, текст скорее всего испорчен
        if suspicious_count > valid_count and suspicious_count > 2:
            return False
        
        # Должен быть хотя бы один валидный символ
        return valid_count > 0
    
    def parse_file(self, file_path: Path) -> List[Tuple[str, str]]:
        """
        Автоматически определяет тип файла и парсит его
        """
        extension = file_path.suffix.lower()
        
        if extension == ".txt":
            return self.parse_txt_file(file_path)
        elif extension == ".pdf":
            return self.parse_pdf_file(file_path)
        elif extension == ".docx":
            return self.parse_docx_file(file_path)
        elif extension in [".xlsx", ".xls"]:
            return self.parse_excel_file(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {extension}")
    
    def build_glossary_dict(
        self, 
        source_lang: str, 
        target_lang: str = "en"
    ) -> Dict[str, Dict[str, str]]:
        """
        Строит словарь глоссария из всех файлов в папке языка
        
        Returns:
            Dict в формате: {
                "source_term": {
                    "source": "исходный термин",
                    "target": "перевод",
                    "abbreviation": "аббревиатура если есть"
                }
            }
        """
        glossary = {}
        lang_dir = self.glossary_dir / source_lang
        
        if not lang_dir.exists():
            print(f"⚠️  Папка {lang_dir} не найдена")
            return glossary
        
        # Обрабатываем все файлы в папке
        for file_path in lang_dir.iterdir():
            if file_path.is_file():
                try:
                    print(f"📄 Обработка файла: {file_path.name}")
                    terms = self.parse_file(file_path)
                    
                    for source_term, target_term in terms:
                        # Извлекаем аббревиатуры из скобок
                        source_abbr = None
                        target_abbr = None
                        
                        # Ищем аббревиатуры в скобках
                        source_match = re.search(r'\(([^)]+)\)', source_term)
                        if source_match:
                            source_abbr = source_match.group(1)
                            source_term = re.sub(r'\s*\([^)]+\)', '', source_term).strip()
                        
                        target_match = re.search(r'\(([^)]+)\)', target_term)
                        if target_match:
                            target_abbr = target_match.group(1)
                            target_term = re.sub(r'\s*\([^)]+\)', '', target_term).strip()
                        
                        # Проверяем, что термины не содержат только подозрительные символы
                        # (результат неправильной кодировки)
                        if self._has_valid_text(source_term) and self._has_valid_text(target_term):
                            # Используем исходный термин как ключ (в нижнем регистре для поиска)
                            key = source_term.lower()
                            
                            glossary[key] = {
                                "source": source_term,
                                "target": target_term,
                                "source_abbr": source_abbr,
                                "target_abbr": target_abbr
                            }
                        # Если текст невалидный, пропускаем его (не добавляем в глоссарий)
                    
                    print(f"✅ Извлечено {len(terms)} терминов из {file_path.name}")
                except Exception as e:
                    print(f"❌ Ошибка при обработке {file_path.name}: {str(e)}")
                    continue
        
        return glossary
    
    def save_glossary_json(
        self, 
        source_lang: str, 
        output_path: Optional[Path] = None,
        target_lang: str = "en"
    ) -> Path:
        """
        Сохраняет глоссарий в JSON файл
        """
        glossary = self.build_glossary_dict(source_lang, target_lang)
        
        if output_path is None:
            output_path = self.glossary_dir / f"glossary_{source_lang}_to_{target_lang}.json"
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(glossary, f, ensure_ascii=False, indent=2)
        
        print(f"💾 Глоссарий сохранен в {output_path}")
        print(f"📊 Всего терминов: {len(glossary)}")
        
        return output_path
    
    def load_glossary_json(self, json_path: Path) -> Dict[str, Dict[str, str]]:
        """
        Загружает глоссарий из JSON файла
        """
        with open(json_path, "r", encoding="utf-8") as f:
            return json.load(f)


def build_all_glossaries():
    """
    Создает JSON файлы для всех языков глоссария
    """
    parser = GlossaryParser()
    
    languages = ["russian", "arabic", "chinise"]
    
    for lang in languages:
        print(f"\n{'='*50}")
        print(f"Обработка глоссария для языка: {lang}")
        print(f"{'='*50}")
        
        try:
            parser.save_glossary_json(lang)
        except Exception as e:
            print(f"❌ Ошибка при обработке {lang}: {str(e)}")


if __name__ == "__main__":
    # Запуск парсинга всех глоссариев
    build_all_glossaries()

