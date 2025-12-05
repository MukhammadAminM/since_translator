from pathlib import Path
from datetime import datetime
from typing import Literal, Optional
import uuid
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Рендерер LaTeX формул
try:
    from services.latex_renderer import LaTeXRenderer
    LATEX_RENDERER_AVAILABLE = True
except ImportError:
    LATEX_RENDERER_AVAILABLE = False


class DocxGenerator:
    """
    Сервис для генерации .docx файлов с переведенным текстом
    """
    
    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Инициализация рендерера LaTeX
        if LATEX_RENDERER_AVAILABLE:
            try:
                self.latex_renderer = LaTeXRenderer()
            except Exception as e:
                print(f"⚠️  Не удалось инициализировать LaTeXRenderer: {str(e)}")
                self.latex_renderer = None
        else:
            self.latex_renderer = None
    
    def create_docx(
        self,
        translated_text: str,
        source_lang: Literal["ru", "ar", "zh"],
        model: Literal["general", "engineering", "academic", "scientific"],
        original_filename: Optional[str] = None,
        original_text: Optional[str] = None,
        page_images: Optional[dict[int, str]] = None
    ) -> str:
        """
        Создает .docx файл с переведенным текстом
        
        Args:
            translated_text: Переведенный текст
            source_lang: Исходный язык
            model: Модель перевода
            original_filename: Имя оригинального файла (если есть)
            original_text: Оригинальный текст (для сохранения структуры)
        
        Returns:
            Имя созданного файла
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx не установлен. Установите: pip install python-docx")
        
        # Создаем новый документ
        doc = Document()
        
        # Настройка страницы
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # Настройка стилей по умолчанию
        self._setup_default_styles(doc)
        
        # Заголовок документа
        self._add_title(doc, original_filename)
        
        # Метаинформация в таблице
        self._add_metadata_table(doc, source_lang, model, original_filename)
        
        # Разделитель
        self._add_separator(doc)
        
        # Основной текст перевода с сохранением структуры
        self._add_translated_content(doc, translated_text, original_text, page_images)
        
        # Сохраняем файл
        filename = self._generate_filename(source_lang, model, original_filename)
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        return filename
    
    def _setup_default_styles(self, doc: Document):
        """Настраивает стили документа"""
        # Стиль Normal
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(6)
        
        # Стиль заголовков
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = 'Calibri'
            heading_font.bold = True
            if i == 1:
                heading_font.size = Pt(16)
                heading_font.color.rgb = RGBColor(31, 78, 121)  # Темно-синий
            elif i == 2:
                heading_font.size = Pt(14)
                heading_font.color.rgb = RGBColor(47, 84, 150)  # Синий
            else:
                heading_font.size = Pt(12)
                heading_font.color.rgb = RGBColor(68, 114, 196)  # Светло-синий
    
    def _add_title(self, doc: Document, original_filename: Optional[str] = None):
        """Добавляет заголовок документа"""
        title_text = "Translation Document"
        if original_filename:
            try:
                # Безопасное извлечение имени файла
                filename_stem = Path(original_filename).stem
                if filename_stem:
                    title_text = f"Translation: {filename_stem}"
            except Exception:
                # Если не удалось обработать имя файла, используем по умолчанию
                pass
        
        title = doc.add_heading(title_text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем подзаголовок
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Generated by Since Translator")
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
        subtitle_run.italic = True
        
        doc.add_paragraph()  # Пустая строка
    
    def _add_metadata_table(self, doc: Document, source_lang: str, model: str, original_filename: Optional[str]):
        """Добавляет таблицу с метаданными"""
        try:
            table = doc.add_table(rows=0, cols=2)
            # Пробуем установить стиль, если не доступен - используем без стиля
            try:
                table.style = 'Light Grid Accent 1'
            except Exception:
                # Если стиль не доступен, используем базовый
                pass
            table.autofit = True
            
            # Настройка ширины колонок
            for col in table.columns:
                col.width = Cm(5)
            
            # Данные для таблицы
            metadata = [
                ("Source Language", source_lang.upper() if source_lang else "Unknown"),
                ("Target Language", "English (EN)"),
                ("Translation Model", model.upper() if model else "Unknown"),
                ("Generated", datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ]
            
            if original_filename:
                # Безопасная обработка имени файла
                try:
                    safe_filename = str(original_filename)[:100]  # Ограничиваем длину
                    metadata.append(("Original File", safe_filename))
                except Exception:
                    pass
        except Exception as e:
            # Если не удалось создать таблицу, создаем простой параграф
            para = doc.add_paragraph()
            para.add_run(f"Source Language: {source_lang.upper() if source_lang else 'Unknown'}").bold = True
            para.add_run(f"\nTarget Language: English (EN)")
            para.add_run(f"\nTranslation Model: {model.upper() if model else 'Unknown'}")
            para.add_run(f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            if original_filename:
                try:
                    para.add_run(f"\nOriginal File: {str(original_filename)[:100]}")
                except Exception:
                    pass
            return
        
        # Добавляем строки
        for label, value in metadata:
            row = table.add_row()
            # Левая колонка (метка)
            label_cell = row.cells[0]
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.color.rgb = RGBColor(68, 114, 196)
            label_para.paragraph_format.space_after = Pt(0)
            
            # Правая колонка (значение)
            value_cell = row.cells[1]
            value_para = value_cell.paragraphs[0]
            value_para.add_run(value)
            value_para.paragraph_format.space_after = Pt(0)
        
        doc.add_paragraph()  # Пустая строка после таблицы
    
    def _add_separator(self, doc: Document):
        """Добавляет визуальный разделитель"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("─" * 60)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(10)
        doc.add_paragraph()  # Пустая строка
    
    def _add_translated_content(self, doc: Document, translated_text: str, original_text: Optional[str] = None, page_images: Optional[dict[int, str]] = None):
        """Добавляет переведенный контент с сохранением структуры и вставкой изображений"""
        if not translated_text or not translated_text.strip():
            # Если текст пустой, добавляем сообщение
            para = doc.add_paragraph("No translated content available.")
            return
        
        # Заголовок раздела
        heading = doc.add_heading('Translated Content', level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        
        # Разбиваем текст на параграфы
        try:
            paragraphs = self._split_into_paragraphs(translated_text, original_text)
        except Exception:
            # Если не удалось разбить, добавляем весь текст как один параграф
            paragraphs = [translated_text]
        
        # Добавляем каждый параграф
        for para_text in paragraphs:
            if not para_text or not para_text.strip():
                # Пустой параграф для разделения
                doc.add_paragraph()
                continue
            
            try:
                # Проверяем, есть ли плейсхолдеры для изображений
                if page_images and '__IMAGE_PAGE_' in para_text:
                    # Заменяем плейсхолдеры на изображения
                    para_text = self._insert_page_images(doc, para_text, page_images)
                    # Если после замены текст пустой, пропускаем
                    if not para_text.strip():
                        continue
                
                # Проверяем, является ли это заголовком
                if self._is_heading(para_text):
                    level = self._get_heading_level(para_text)
                    heading_para = doc.add_heading(para_text.strip('# '), level=level)
                    heading_para.paragraph_format.space_before = Pt(12)
                    heading_para.paragraph_format.space_after = Pt(6)
                else:
                    # Проверяем, содержит ли параграф LaTeX формулы
                    if self._contains_latex_formula(para_text):
                        # Обрабатываем параграф с формулами
                        self._add_paragraph_with_formulas(doc, para_text)
                    else:
                        # Обычный параграф
                        para = doc.add_paragraph(para_text.strip())
                        para.paragraph_format.first_line_indent = Cm(0.5)  # Отступ первой строки
                        para.paragraph_format.line_spacing = 1.15
            except Exception as e:
                # Если не удалось добавить параграф, пробуем простой вариант
                try:
                    doc.add_paragraph(para_text.strip())
                except Exception:
                    # Если и это не работает, пропускаем параграф
                    continue
    
    def _split_into_paragraphs(self, translated_text: str, original_text: Optional[str] = None) -> list:
        """Разбивает текст на параграфы, сохраняя структуру"""
        # Разбиваем по двойным переносам строк
        paragraphs = re.split(r'\n\s*\n', translated_text)
        
        # Очищаем параграфы
        cleaned = []
        for para in paragraphs:
            para = para.strip()
            if para:
                # Разбиваем длинные параграфы на предложения для лучшей читаемости
                if len(para) > 500:
                    # Разбиваем по точкам, но сохраняем структуру
                    sentences = re.split(r'(?<=[.!?])\s+', para)
                    current_para = ""
                    for sentence in sentences:
                        if len(current_para) + len(sentence) < 400:
                            current_para += sentence + " "
                        else:
                            if current_para:
                                cleaned.append(current_para.strip())
                            current_para = sentence + " "
                    if current_para:
                        cleaned.append(current_para.strip())
                else:
                    cleaned.append(para)
        
        return cleaned if cleaned else [translated_text]
    
    def _is_heading(self, text: str) -> bool:
        """Проверяет, является ли текст заголовком"""
        text = text.strip()
        # Проверяем на markdown заголовки
        if text.startswith('#'):
            return True
        # Проверяем на короткие строки в верхнем регистре
        if len(text) < 100 and text.isupper() and len(text.split()) < 10:
            return True
        return False
    
    def _get_heading_level(self, text: str) -> int:
        """Определяет уровень заголовка"""
        text = text.strip()
        if text.startswith('#'):
            level = len(text) - len(text.lstrip('#'))
            return min(level, 3)  # Максимум 3 уровень
        return 2  # По умолчанию уровень 2
    
    def _contains_latex_formula(self, text: str) -> bool:
        """Проверяет, содержит ли текст LaTeX формулы"""
        # Проверяем на \[ ... \] или \( ... \)
        return bool(re.search(r'\\\[.*?\\\]|\\\(.*?\\\)', text, re.DOTALL))
    
    def _add_paragraph_with_formulas(self, doc: Document, text: str):
        """Добавляет параграф с LaTeX формулами, рендеря их в изображения"""
        # Находим все LaTeX формулы
        latex_pattern = r'(\\\[.*?\\\]|\\\(.*?\\\))'
        parts = re.split(latex_pattern, text)
        
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0.5)
        para.paragraph_format.line_spacing = 1.15
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for part in parts:
            if not part.strip():
                continue
            
            # Проверяем, является ли это LaTeX формулой
            if re.match(r'\\\[.*?\\\]|\\\(.*?\\\)', part, re.DOTALL):
                # Рендерим формулу в изображение
                if self.latex_renderer and self.latex_renderer.available:
                    formula_image = self.latex_renderer.render_latex_to_image(part)
                    if formula_image:
                        # Вставляем изображение в параграф
                        run = para.add_run()
                        run.add_picture(formula_image, width=Inches(4))  # Ширина 4 дюйма
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Если не удалось отрендерить, вставляем как текст
                        run = para.add_run(f"[Formula: {part[:50]}...]")
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(128, 128, 128)
                else:
                    # Если рендерер недоступен, вставляем как текст
                    run = para.add_run(f"[Formula: {part[:50]}...]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)
            else:
                # Обычный текст
                para.add_run(part)
    
    def _insert_page_images(self, doc: Document, text: str, page_images: dict[int, str]) -> str:
        """
        Вставляет изображения страниц вместо плейсхолдеров
        
        Args:
            doc: Word документ
            text: Текст с плейсхолдерами __IMAGE_PAGE_{page}__ или __IMAGE_PAGE_{page}_LINE_{line}__
            page_images: Словарь {номер_страницы: путь_к_изображению}
        
        Returns:
            Текст без плейсхолдеров (они заменены на пустую строку, изображения вставлены)
        """
        import re
        from pathlib import Path
        
        # Находим все плейсхолдеры изображений
        # Поддерживаем два формата: __IMAGE_PAGE_{page}__ и __IMAGE_PAGE_{page}_LINE_{line}__
        pattern = r'__IMAGE_PAGE_(\d+)(?:_LINE_\d+)?__'
        matches = list(re.finditer(pattern, text))
        
        if not matches:
            return text
        
        # Вставляем изображения в обратном порядке, чтобы не сбить позиции
        for match in reversed(matches):
            page_num = int(match.group(1))
            if page_num in page_images:
                image_path = Path(page_images[page_num])
                if image_path.exists():
                    try:
                        # Создаем параграф для изображения
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        # Вставляем изображение с ограничением размера (ширина 6 дюймов для страницы)
                        run.add_picture(str(image_path), width=Inches(6))
                        print(f"   📷 Вставлено изображение страницы {page_num} в Word документ")
                    except Exception as e:
                        print(f"   ⚠️  Не удалось вставить изображение страницы {page_num}: {e}")
                else:
                    print(f"   ⚠️  Изображение страницы {page_num} не найдено: {image_path}")
            else:
                print(f"   ⚠️  Номер страницы {page_num} не найден в словаре изображений")
            
            # Удаляем плейсхолдер из текста
            text = text[:match.start()] + text[match.end():]
        
        return text
    
    def _generate_filename(
        self,
        source_lang: str,
        model: str,
        original_filename: Optional[str] = None
    ) -> str:
        """
        Генерирует уникальное имя файла
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        
        if original_filename:
            base_name = Path(original_filename).stem
            return f"{base_name}_translated_{timestamp}_{unique_id}.docx"
        else:
            return f"translation_{source_lang}_to_en_{model}_{timestamp}_{unique_id}.docx"


