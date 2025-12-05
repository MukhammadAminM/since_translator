"""
Модуль для сборки DOCX документов с формулами
Поддерживает два режима:
- PNG: формулы как изображения (быстрый)
- OMML: нативные формулы Word (продвинутый)
"""
import logging
from pathlib import Path
from typing import Literal, Optional, Dict, List
from enum import Enum
import re

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx не установлен. Установите: pip install python-docx")

try:
    from services.latex_renderer import LaTeXRenderer
    LATEX_RENDERER_AVAILABLE = True
except ImportError:
    LATEX_RENDERER_AVAILABLE = False
    logger.warning("LaTeXRenderer недоступен")

try:
    import lxml.etree as ET
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False
    logger.warning("lxml не установлен. OMML режим будет недоступен. Установите: pip install lxml")


class FormulaMode(Enum):
    """Режимы отображения формул"""
    PNG = "png"  # Формулы как изображения (быстрый)
    OMML = "omml"  # Нативные формулы Word (продвинутый)


class DocumentBuilder:
    """
    Класс для сборки DOCX документов с формулами
    
    Поддерживает:
    - PNG режим: LaTeX → PNG изображения
    - OMML режим: LaTeX → MathML → OMML (нативные формулы Word)
    """
    
    def __init__(
        self,
        output_dir: str = "outputs",
        formula_mode: FormulaMode = FormulaMode.PNG
    ):
        """
        Args:
            output_dir: Директория для сохранения файлов
            formula_mode: Режим отображения формул
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx не установлен")
        
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.formula_mode = formula_mode
        
        # Инициализация рендерера LaTeX (для PNG режима)
        if formula_mode == FormulaMode.PNG and LATEX_RENDERER_AVAILABLE:
            try:
                self.latex_renderer = LaTeXRenderer()
                logger.info("LaTeXRenderer инициализирован для PNG режима")
            except Exception as e:
                logger.warning(f"Не удалось инициализировать LaTeXRenderer: {e}")
                self.latex_renderer = None
        else:
            self.latex_renderer = None
        
        # Проверка доступности OMML
        if formula_mode == FormulaMode.OMML and not LXML_AVAILABLE:
            logger.warning("lxml не установлен, переключаемся на PNG режим")
            self.formula_mode = FormulaMode.PNG
    
    def build(
        self,
        translated_text: str,
        formulas: Dict[int, str],
        recognized_formulas: Optional[Dict[int, 'RecognizedFormula']] = None,
        source_lang: Literal["ru", "ar", "zh"] = "ru",
        model: Literal["general", "engineering", "academic", "scientific"] = "general",
        original_filename: Optional[str] = None,
        page_images: Optional[Dict[int, Path]] = None
    ) -> Path:
        """
        Собирает DOCX документ
        
        Args:
            translated_text: Переведенный текст с маркерами формул
            formulas: Словарь формул {индекс: оригинальная_формула}
            recognized_formulas: Словарь распознанных формул {индекс: RecognizedFormula}
            source_lang: Исходный язык
            model: Модель перевода
            original_filename: Имя оригинального файла
            page_images: Изображения страниц для вставки
        
        Returns:
            Путь к созданному файлу
        """
        logger.info(f"Начало сборки DOCX (режим формул: {self.formula_mode.value})")
        
        # Создаем документ
        doc = Document()
        
        # Настройка страницы
        self._setup_page(doc)
        
        # Заголовок и метаданные
        self._add_header(doc, original_filename, source_lang, model)
        
        # Добавляем контент
        self._add_content(
            doc,
            translated_text,
            formulas,
            recognized_formulas,
            page_images
        )
        
        # Сохраняем файл
        filename = self._generate_filename(source_lang, model, original_filename)
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        logger.info(f"DOCX файл создан: {filepath}")
        
        return filepath
    
    def _setup_page(self, doc: Document):
        """Настраивает параметры страницы"""
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    def _add_header(
        self,
        doc: Document,
        original_filename: Optional[str],
        source_lang: str,
        model: str
    ):
        """Добавляет заголовок и метаданные"""
        # Заголовок
        if original_filename:
            title = doc.add_heading(f"Перевод: {original_filename}", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Метаданные
        from datetime import datetime
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Исходный язык: {source_lang.upper()}")
        meta_para.add_run(" | ")
        meta_para.add_run(f"Модель: {model}")
        meta_para.add_run(" | ")
        meta_para.add_run(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Разделитель
        doc.add_paragraph("─" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_content(
        self,
        doc: Document,
        translated_text: str,
        formulas: Dict[int, str],
        recognized_formulas: Optional[Dict[int, 'RecognizedFormula']],
        page_images: Optional[Dict[int, Path]]
    ):
        """Добавляет основной контент с формулами"""
        # Разбиваем текст на параграфы
        paragraphs = translated_text.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            # Проверяем наличие маркеров формул
            if '<<<FORMULA_' in para_text:
                self._add_paragraph_with_formulas(
                    doc,
                    para_text,
                    formulas,
                    recognized_formulas
                )
            else:
                # Обычный параграф
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.first_line_indent = Cm(0.5)
                para.paragraph_format.line_spacing = 1.15
        
        # Вставляем изображения страниц, если есть
        if page_images:
            self._insert_page_images(doc, translated_text, page_images)
    
    def _add_paragraph_with_formulas(
        self,
        doc: Document,
        text: str,
        formulas: Dict[int, str],
        recognized_formulas: Optional[Dict[int, 'RecognizedFormula']]
    ):
        """Добавляет параграф с формулами"""
        # Находим все маркеры формул
        pattern = r'<<<FORMULA_(\d+)>>>'
        parts = re.split(pattern, text)
        
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0.5)
        para.paragraph_format.line_spacing = 1.15
        
        i = 0
        while i < len(parts):
            part = parts[i].strip()
            
            if part:
                # Обычный текст
                para.add_run(part)
            
            # Проверяем, есть ли следующий элемент (индекс формулы)
            if i + 1 < len(parts):
                formula_index = int(parts[i + 1])
                
                # Получаем формулу
                if formula_index in formulas:
                    formula_text = formulas[formula_index]
                    
                    # Пробуем получить распознанную формулу
                    recognized = None
                    if recognized_formulas and formula_index in recognized_formulas:
                        recognized = recognized_formulas[formula_index]
                    
                    # Вставляем формулу в зависимости от режима
                    if self.formula_mode == FormulaMode.PNG:
                        self._insert_formula_as_image(para, formula_text, recognized)
                    elif self.formula_mode == FormulaMode.OMML:
                        self._insert_formula_as_omml(para, formula_text, recognized)
                    else:
                        # Fallback: вставляем как текст
                        para.add_run(f"[Formula: {formula_text[:50]}...]")
                
                i += 2
            else:
                i += 1
        
        # Центрируем параграф, если содержит формулы
        if '<<<FORMULA_' in text:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _insert_formula_as_image(
        self,
        para,
        formula_text: str,
        recognized: Optional['RecognizedFormula']
    ):
        """Вставляет формулу как изображение (PNG режим)"""
        # Используем распознанную формулу, если есть
        if recognized and recognized.latex:
            latex = recognized.latex
        else:
            latex = formula_text
        
        # Рендерим в изображение
        if self.latex_renderer:
            image_buf = self.latex_renderer.render_latex_to_image(latex)
            if image_buf:
                run = para.add_run()
                run.add_picture(image_buf, width=Inches(4))
                logger.debug(f"Формула вставлена как изображение: {latex[:50]}...")
                return
        
        # Fallback: вставляем как текст
        para.add_run(f"[Formula: {formula_text[:50]}...]")
        logger.warning(f"Не удалось отрендерить формулу: {formula_text[:50]}...")
    
    def _insert_formula_as_omml(
        self,
        para,
        formula_text: str,
        recognized: Optional['RecognizedFormula']
    ):
        """Вставляет формулу как OMML (OMML режим)"""
        if not LXML_AVAILABLE:
            logger.warning("lxml недоступен, используем PNG режим")
            self._insert_formula_as_image(para, formula_text, recognized)
            return
        
        # Получаем MathML
        mathml = None
        if recognized and recognized.mathml:
            mathml = recognized.mathml
        else:
            # Пробуем конвертировать LaTeX в MathML
            # Это упрощенная версия - в реальности нужна библиотека для конвертации
            logger.warning("MathML недоступен, используем PNG режим")
            self._insert_formula_as_image(para, formula_text, recognized)
            return
        
        # Конвертируем MathML в OMML и вставляем
        try:
            omml = self._mathml_to_omml(mathml)
            self._insert_omml(para, omml)
            logger.debug(f"Формула вставлена как OMML: {formula_text[:50]}...")
        except Exception as e:
            logger.error(f"Ошибка при конвертации в OMML: {e}")
            # Fallback на PNG
            self._insert_formula_as_image(para, formula_text, recognized)
    
    def _mathml_to_omml(self, mathml: str) -> str:
        """
        Конвертирует MathML в OMML
        
        Args:
            mathml: MathML строка
        
        Returns:
            OMML XML строка
        """
        # Это упрощенная версия
        # В реальности нужна более сложная конвертация
        # Можно использовать библиотеку pymathml или написать свой конвертер
        
        # Пока возвращаем упрощенный OMML
        # В продакшене нужно использовать специализированную библиотеку
        logger.warning("Конвертация MathML → OMML требует дополнительной реализации")
        raise NotImplementedError("MathML → OMML конвертация требует дополнительной реализации")
    
    def _insert_omml(self, para, omml: str):
        """Вставляет OMML в параграф"""
        # Создаем элемент для формулы
        math_element = OxmlElement('m:oMathPara')
        math_element.set(qn('xmlns:m'), 'http://schemas.openxmlformats.org/officeDocument/2006/math')
        
        # Парсим OMML и вставляем
        # Это упрощенная версия - в реальности нужен полный парсер OMML
        para._element.append(math_element)
    
    def _insert_page_images(
        self,
        doc: Document,
        text: str,
        page_images: Dict[int, Path]
    ):
        """Вставляет изображения страниц вместо плейсхолдеров"""
        pattern = r'__IMAGE_PAGE_(\d+)__'
        matches = list(re.finditer(pattern, text))
        
        for match in reversed(matches):
            page_num = int(match.group(1))
            if page_num in page_images:
                image_path = page_images[page_num]
                if image_path.exists():
                    try:
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(str(image_path), width=Inches(6))
                        logger.debug(f"Вставлено изображение страницы {page_num}")
                    except Exception as e:
                        logger.error(f"Ошибка при вставке изображения страницы {page_num}: {e}")
    
    def _generate_filename(
        self,
        source_lang: str,
        model: str,
        original_filename: Optional[str]
    ) -> str:
        """Генерирует имя файла"""
        from datetime import datetime
        import uuid
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        
        if original_filename:
            base_name = Path(original_filename).stem
            return f"{base_name}_translated_{timestamp}_{unique_id}.docx"
        else:
            return f"translated_{source_lang}_{model}_{timestamp}_{unique_id}.docx"

